"""The rich-text whitelist: editor HTML -> blocks -> docx."""
import io

from docx import Document

from template_builder.richtext import add_blocks, blocks_from_text, parse_html


def test_paragraphs_and_inline_formatting():
    blocks = parse_html("<p>The <b>aggregate cap</b> governs.</p><p>Second para.</p>")
    assert [b.kind for b in blocks] == ["p", "p"]
    runs = blocks[0].runs
    assert [(r.text, r.bold) for r in runs] == [
        ("The ", False), ("aggregate cap", True), (" governs.", False)]
    # the boundary spaces around the bold run survive
    assert blocks[0].text() == "The aggregate cap governs."


def test_nested_and_synonym_tags():
    blocks = parse_html("<div><strong><em>both</em></strong> and <u>under</u></div>")
    styles = [(r.text, r.bold, r.italic, r.underline) for r in blocks[0].runs]
    assert styles == [("both", True, True, False), (" and ", False, False, False),
                      ("under", False, False, True)]


def test_lists_bullet_and_numbered():
    blocks = parse_html(
        "<p>Intro:</p><ul><li>first</li><li><b>second</b></li></ul>"
        "<ol><li>step one</li></ol>")
    assert [b.kind for b in blocks] == ["p", "bullet", "bullet", "number"]
    assert blocks[2].runs[0].bold is True
    assert blocks[3].text() == "step one"


def test_junk_is_stripped_but_content_kept():
    html = ('<span style="font-weight:700" class="x">kept</span>'
            '<script>alert(1)</script><font color="red"> also kept</font>'
            '<p>&nbsp;&nbsp;</p>')
    blocks = parse_html(html)
    assert len(blocks) == 1
    assert blocks[0].text() == "kept also kept"
    assert all(not r.bold for r in blocks[0].runs)  # style attrs carry no meaning


def test_br_and_divs_split_paragraphs():
    blocks = parse_html("first line<br>second<div>third</div>")
    assert [b.text() for b in blocks] == ["first line", "second", "third"]


def test_malformed_html_never_raises():
    assert parse_html("<b>unclosed <li>stray</ul></i>") != []
    assert parse_html("") == []
    assert parse_html("   ") == []


def test_blocks_from_text_fallback():
    blocks = blocks_from_text("Para one.\n\nPara two.")
    assert [b.text() for b in blocks] == ["Para one.", "Para two."]
    assert all(b.kind == "p" for b in blocks)


def test_docx_emission_styles_and_runs():
    doc = Document()
    add_blocks(doc, parse_html(
        "<p>Plain with <b>bold</b>.</p><ul><li>item</li></ul><ol><li>num</li></ol>"))
    paragraphs = doc.paragraphs
    assert [p.style.name for p in paragraphs] == ["Normal", "List Bullet", "List Number"]
    assert [r.bold for r in paragraphs[0].runs] == [None, True, None]
    # round-trips through a real .docx file
    buffer = io.BytesIO()
    doc.save(buffer)
    reloaded = Document(io.BytesIO(buffer.getvalue()))
    assert reloaded.paragraphs[1].style.name == "List Bullet"


def test_export_endpoint_uses_rich_html(tmp_path, template_dict):
    import json

    from fastapi.testclient import TestClient

    from template_builder.server import create_app
    (tmp_path / "nda.json").write_text(json.dumps(template_dict))
    client = TestClient(create_app(str(tmp_path)))
    res = client.post("/api/export-docx", json={
        "title": "Rich NDA",
        "clauses": [{
            "id": "a", "number": 1, "heading": "Obligations",
            "text": "fallback text",
            "html": "<p>Must <b>never</b> disclose:</p><ul><li>plans</li><li>pricing</li></ul>",
        }],
    })
    assert res.status_code == 200
    doc = Document(io.BytesIO(res.content))
    texts = [(p.text, p.style.name) for p in doc.paragraphs]
    assert ("Must never disclose:", "Normal") in texts
    assert ("plans", "List Bullet") in texts and ("pricing", "List Bullet") in texts
    bold_runs = [r.text for p in doc.paragraphs for r in p.runs if r.bold]
    assert bold_runs == ["never"]
