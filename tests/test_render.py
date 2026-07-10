"""Behavioural tests for template_builder.render.

Complements the anchor tests in test_core.py (happy path, determinism,
numbering shifts, missing/mistyped answers, refs to excluded clauses).
Here: placeholder syntax tolerance, value formatting, variant selection
order, plan numbering, error aggregation, and DOCX output.
"""
import pytest
from docx import Document

from template_builder.model import template_from_dict
from template_builder.render import (
    RenderError,
    check_answers,
    plan,
    render_docx,
    render_markdown,
    select_variant,
)


# ---------------------------------------------------------------------------
# check_answers
# ---------------------------------------------------------------------------

def test_check_answers_returns_no_problems_for_valid_answers(template, answers):
    assert check_answers(template, answers) == []


def test_extra_answer_key_raises_render_error(template, answers):
    answers["shoe_size"] = 11
    with pytest.raises(RenderError) as exc:
        render_markdown(template, answers)
    message = str(exc.value)
    assert "shoe_size" in message
    assert "does not match any template variable" in message


def test_check_answers_rejects_boolean_for_number_type(template, answers):
    # bool is a subclass of int, so this guards against True sneaking in as 1.
    answers["term_years"] = True
    problems = check_answers(template, answers)
    assert len(problems) == 1
    assert "term_years" in problems[0] and "must be a number" in problems[0]


# ---------------------------------------------------------------------------
# Placeholder syntax
# ---------------------------------------------------------------------------

def test_placeholders_tolerate_interior_whitespace(template_dict, answers):
    clauses = {c["id"]: c for c in template_dict["clauses"]}
    clauses["parties"]["variants"][0]["text"] = (
        "This Agreement is made on {{ effective_date }} between "
        "{{  party_1  }} and {{party_2 }}."
    )
    clauses["obligations"]["variants"][0]["text"] = (
        "Each party shall protect Confidential Information as defined in "
        "{{ ref:definitions }}."
    )
    doc = render_markdown(template_from_dict(template_dict), answers)
    assert "made on 1 July 2026 between Acme Corp and Blue Ridge LLC." in doc
    assert "as defined in clause 2." in doc
    assert "{{" not in doc and "}}" not in doc


def test_malformed_placeholder_is_reported_not_silently_kept(template_dict, answers):
    # A space after "ref:" is outside the placeholder grammar; rendering must
    # fail loudly rather than shipping literal braces in a contract.
    template_dict["clauses"][0]["variants"][0]["text"] += " See {{ref: definitions}}."
    with pytest.raises(RenderError) as exc:
        render_markdown(template_from_dict(template_dict), answers)
    message = str(exc.value)
    assert "'parties'" in message
    assert "unresolved placeholder syntax" in message


# ---------------------------------------------------------------------------
# Value formatting (via rendering, exercising _format_value)
# ---------------------------------------------------------------------------

def test_boolean_answers_render_as_lowercase_true_false(template_dict, answers):
    template_dict["clauses"][0]["variants"][0]["text"] += " Mutual: {{is_mutual}}."
    template = template_from_dict(template_dict)

    doc = render_markdown(template, answers)  # is_mutual is True
    assert "Mutual: true." in doc
    assert "Mutual: True." not in doc

    answers["is_mutual"] = False
    doc = render_markdown(template, answers)
    assert "Mutual: false." in doc


def test_integral_float_renders_without_decimal_point(template, answers):
    answers["term_years"] = 3.0
    doc = render_markdown(template, answers)
    assert "continues for 3 years" in doc
    assert "3.0" not in doc


def test_non_integral_float_renders_with_its_decimals(template, answers):
    answers["term_years"] = 2.5
    doc = render_markdown(template, answers)
    assert "continues for 2.5 years" in doc


# ---------------------------------------------------------------------------
# select_variant
# ---------------------------------------------------------------------------

def test_select_variant_first_match_wins_when_multiple_conditions_true(template_dict, answers):
    gov = next(c for c in template_dict["clauses"] if c["id"] == "governing-law")
    gov["variants"].insert(0, {
        "id": "broad",
        "when": 'governing_law in ["New York", "Delaware"]',
        "provenance": [],
        "text": "This Agreement is governed by the chosen state's laws.",
    })
    template = template_from_dict(template_dict)
    # answers["governing_law"] == "New York": both "broad" and "ny" conditions hold.
    chosen = select_variant(template.clause("governing-law"), answers)
    assert chosen.id == "broad"

    doc = render_markdown(template, answers)
    assert "chosen state's laws" in doc
    assert "State of New York" not in doc


def test_select_variant_falls_through_to_default_and_returns_none_without_one(template, answers):
    obligations = template.clause("obligations")
    answers["is_mutual"] = False
    assert select_variant(obligations, answers).id == "one-way"

    # Strip the default: with is_mutual False nothing matches at all.
    obligations.variants = [v for v in obligations.variants if v.when is not None]
    assert select_variant(obligations, answers) is None


def test_clause_with_no_matching_variant_fails_the_plan(template_dict, answers):
    obligations = next(c for c in template_dict["clauses"] if c["id"] == "obligations")
    obligations["variants"] = [v for v in obligations["variants"] if v["when"] is not None]
    answers["is_mutual"] = False
    with pytest.raises(RenderError) as exc:
        plan(template_from_dict(template_dict), answers)
    message = str(exc.value)
    assert "'obligations'" in message and "no variant matches" in message


# ---------------------------------------------------------------------------
# plan
# ---------------------------------------------------------------------------

def test_plan_numbers_only_included_clauses_contiguously(template, answers):
    planned = plan(template, answers)  # include_non_solicit is False
    assert [p.clause.id for p in planned] == [
        "parties", "definitions", "obligations", "term", "governing-law",
    ]
    assert [p.number for p in planned] == [1, 2, 3, 4, 5]

    answers["include_non_solicit"] = True
    planned = plan(template, answers)
    assert [p.clause.id for p in planned] == [
        "parties", "definitions", "obligations", "non-solicit", "term", "governing-law",
    ]
    assert [p.number for p in planned] == [1, 2, 3, 4, 5, 6]


def test_plan_fills_in_rendered_text(template, answers):
    planned = plan(template, answers)
    by_id = {p.clause.id: p for p in planned}
    assert "Acme Corp" in by_id["parties"].rendered_text
    assert "as defined in clause 2" in by_id["obligations"].rendered_text
    assert "{{" not in by_id["term"].rendered_text


# ---------------------------------------------------------------------------
# Substitution errors
# ---------------------------------------------------------------------------

def test_unknown_variable_in_text_names_the_clause(template_dict, answers):
    template_dict["clauses"][4]["variants"][0]["text"] += " Renews for {{renewal_years}} years."
    with pytest.raises(RenderError) as exc:
        render_markdown(template_from_dict(template_dict), answers)
    message = str(exc.value)
    assert "'term'" in message
    assert "unknown variable 'renewal_years'" in message


def test_ref_to_nonexistent_clause_says_not_in_the_template(template_dict, answers):
    template_dict["clauses"][0]["variants"][0]["text"] += " See {{ref:arbitration}}."
    with pytest.raises(RenderError) as exc:
        render_markdown(template_from_dict(template_dict), answers)
    message = str(exc.value)
    assert "'parties'" in message and "'arbitration'" in message
    assert "not in the template" in message
    assert "excluded under these answers" not in message


def test_render_error_aggregates_problems_across_clauses(template_dict, answers):
    # Two independent defects in two different clauses -> one raise, both listed.
    template_dict["clauses"][0]["variants"][0]["text"] += " Signed by {{witness_name}}."
    template_dict["clauses"][4]["variants"][0]["text"] += " See {{ref:arbitration}}."
    with pytest.raises(RenderError) as exc:
        render_markdown(template_from_dict(template_dict), answers)
    problems = exc.value.problems
    assert len(problems) >= 2
    assert any("witness_name" in p for p in problems)
    assert any("arbitration" in p for p in problems)
    # The exception text is the joined problem list, one bullet per problem.
    assert str(exc.value).count("\n  - ") == len(problems)


# ---------------------------------------------------------------------------
# render_markdown / render_docx titles and DOCX output
# ---------------------------------------------------------------------------

def test_render_markdown_custom_title_overrides_doc_type(template, answers):
    doc = render_markdown(template, answers, title="Project Falcon NDA")
    assert doc.splitlines()[0] == "# Project Falcon NDA"
    assert "# Mutual Non-Disclosure Agreement" not in doc


def test_render_docx_writes_loadable_file_with_headings(template, answers, tmp_path):
    path = tmp_path / "nda.docx"
    render_docx(template, answers, str(path))

    loaded = Document(str(path))
    paragraphs = [(p.style.name, p.text) for p in loaded.paragraphs]
    headings = [text for style, text in paragraphs if style.startswith("Heading")]

    assert ("Heading 1", "Mutual Non-Disclosure Agreement") in paragraphs
    for expected in ["1. Parties", "2. Definitions", "3. Confidentiality Obligations",
                     "4. Term", "5. Governing Law"]:
        assert ("Heading 2", expected) in paragraphs
    assert "4. Non-Solicitation" not in headings  # excluded clause stays out

    body = "\n".join(text for _, text in paragraphs)
    assert "Acme Corp" in body and "Blue Ridge LLC" in body
    assert "as defined in clause 2" in body
    assert "{{" not in body


def test_render_docx_custom_title(template, answers, tmp_path):
    path = tmp_path / "titled.docx"
    render_docx(template, answers, str(path), title="Project Falcon NDA")
    loaded = Document(str(path))
    h1 = [p.text for p in loaded.paragraphs if p.style.name == "Heading 1"]
    assert h1 == ["Project Falcon NDA"]


def test_render_docx_rejects_bad_answers_before_writing(template, answers, tmp_path):
    path = tmp_path / "never.docx"
    del answers["party_1"]
    with pytest.raises(RenderError):
        render_docx(template, answers, str(path))
    assert not path.exists()
