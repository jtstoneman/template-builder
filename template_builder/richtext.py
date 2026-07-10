"""A deliberately tiny rich-text model: contenteditable HTML -> docx.

The document editor supports exactly what .docx contracts need day-to-day —
bold, italic, underline, bullet and numbered lists, paragraphs — and nothing
else. This module is the whitelist that makes that promise hold: it parses
the editor's HTML, keeps only those constructs, and drops every attribute,
style and unknown tag (contenteditable and pasted content produce plenty).

Pure and deterministic: same HTML in, same block structure out, so the
conversion is unit-testable without python-docx.
"""

import re
from dataclasses import dataclass, field
from html.parser import HTMLParser

BOLD_TAGS = {"b", "strong"}
ITALIC_TAGS = {"i", "em"}
UNDERLINE_TAGS = {"u"}
BLOCK_TAGS = {"p", "div", "li"}
LIST_TAGS = {"ul", "ol"}
SKIP_TAGS = {"script", "style"}  # never render their text


@dataclass(slots=True)
class Run:
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False


@dataclass(slots=True)
class Block:
    kind: str  # "p" | "bullet" | "number"
    runs: list[Run] = field(default_factory=list)
    starts_list: bool = False  # first item of a new list: numbering restarts here

    def text(self) -> str:
        return "".join(run.text for run in self.runs)


class _Parser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocks: list[Block] = []
        self._current: Block | None = None
        self._bold = 0
        self._italic = 0
        self._underline = 0
        self._lists: list[str] = []   # stack of "bullet" / "number"
        self._list_opened = False     # the next list item starts a fresh list
        self._skip = 0

    # -- block management --

    def _block_kind(self) -> str:
        return self._lists[-1] if self._lists else "p"

    def _open_block(self, kind: str | None = None) -> Block:
        if self._current is None:
            self._current = Block(kind=kind or self._block_kind())
            if self._current.kind in ("bullet", "number") and self._list_opened:
                self._current.starts_list = True
                self._list_opened = False
        return self._current

    def _close_block(self) -> None:
        if self._current is not None and self._current.text().strip():
            self.blocks.append(self._current)
        self._current = None

    # -- HTMLParser hooks --

    def handle_starttag(self, tag, attrs):
        if tag in SKIP_TAGS:
            self._skip += 1
        elif tag in BOLD_TAGS:
            self._bold += 1
        elif tag in ITALIC_TAGS:
            self._italic += 1
        elif tag in UNDERLINE_TAGS:
            self._underline += 1
        elif tag in LIST_TAGS:
            self._close_block()
            self._lists.append("number" if tag == "ol" else "bullet")
            self._list_opened = True
        elif tag in BLOCK_TAGS:
            self._close_block()
            if tag == "li":
                self._open_block(self._block_kind())
        elif tag == "br":
            self._close_block()
        # every other tag (span, font, a, ...) is transparent: children kept,
        # the tag and all its attributes dropped

    def handle_endtag(self, tag):
        if tag in SKIP_TAGS:
            self._skip = max(0, self._skip - 1)
        elif tag in BOLD_TAGS:
            self._bold = max(0, self._bold - 1)
        elif tag in ITALIC_TAGS:
            self._italic = max(0, self._italic - 1)
        elif tag in UNDERLINE_TAGS:
            self._underline = max(0, self._underline - 1)
        elif tag in LIST_TAGS:
            self._close_block()
            if self._lists:
                self._lists.pop()
        elif tag in BLOCK_TAGS:
            self._close_block()

    def handle_data(self, data):
        if self._skip:
            return
        text = data.replace("\xa0", " ")
        if not text.strip() and self._current is None:
            return  # inter-tag whitespace
        block = self._open_block()
        style = (bool(self._bold), bool(self._italic), bool(self._underline))
        # merge with the previous run when formatting is identical
        if block.runs and (block.runs[-1].bold, block.runs[-1].italic,
                           block.runs[-1].underline) == style:
            block.runs[-1].text += text
        else:
            block.runs.append(Run(text=text, bold=style[0], italic=style[1],
                                  underline=style[2]))


def parse_html(html: str) -> list[Block]:
    """Editor HTML -> whitelisted blocks. Junk-tolerant, never raises on markup."""
    parser = _Parser()
    parser.feed(html)
    parser.close()  # flush buffered trailing text (e.g. an unterminated "M&A")
    parser._close_block()
    blocks = []
    for block in parser.blocks:
        # collapse whitespace runs to single spaces, PRESERVING the boundary
        # space between differently-formatted runs ("the <b>cap</b> applies")
        for run in block.runs:
            run.text = re.sub(r"\s+", " ", run.text)
        runs = [r for r in block.runs if r.text]
        if runs:
            runs[0].text = runs[0].text.lstrip()
            runs[-1].text = runs[-1].text.rstrip()
        runs = [r for r in runs if r.text]
        if runs:
            blocks.append(Block(kind=block.kind, runs=runs,
                                starts_list=block.starts_list))
    return blocks


def blocks_from_text(text: str) -> list[Block]:
    """The plain-text fallback: blank-line-separated paragraphs."""
    return [Block(kind="p", runs=[Run(text=" ".join(p.split()))])
            for p in text.replace("\r\n", "\n").split("\n\n") if p.strip()]


_DOCX_STYLES = {"bullet": "List Bullet", "number": "List Number"}


def _restart_numbering(doc, paragraph) -> None:
    """Give this paragraph a fresh numbering instance so its list starts at 1.

    Word numbers every "List Number" paragraph in one continuous sequence
    unless told otherwise: a second list would silently continue 3., 4. —
    corrupting the enumeration lawyers cross-reference by eye. The fix is a
    new <w:num> pointing at the style's abstract definition with a
    startOverride, attached to the list's first paragraph.
    """
    from docx.oxml.ns import qn

    try:
        numbering = doc.part.numbering_part.element
    except (NotImplementedError, KeyError, AttributeError):
        return  # no numbering definitions to restart against
    style_num_id = None
    for style in doc.styles.element.findall(qn("w:style")):
        if style.get(qn("w:styleId")) == "ListNumber":
            num_ref = style.find(f"{qn('w:pPr')}/{qn('w:numPr')}/{qn('w:numId')}")
            if num_ref is not None:
                style_num_id = num_ref.get(qn("w:val"))
            break
    nums = numbering.findall(qn("w:num"))
    abstract_id = next(
        (n.find(qn("w:abstractNumId")).get(qn("w:val")) for n in nums
         if n.get(qn("w:numId")) == style_num_id
         and n.find(qn("w:abstractNumId")) is not None), None)
    if abstract_id is None:
        return
    new_id = max(int(n.get(qn("w:numId"))) for n in nums) + 1
    num = numbering.makeelement(qn("w:num"), {qn("w:numId"): str(new_id)})
    ref = num.makeelement(qn("w:abstractNumId"), {qn("w:val"): abstract_id})
    override = num.makeelement(qn("w:lvlOverride"), {qn("w:ilvl"): "0"})
    override.append(override.makeelement(qn("w:startOverride"), {qn("w:val"): "1"}))
    num.append(ref)
    num.append(override)
    numbering.append(num)
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = new_id


def add_blocks(doc, blocks: list[Block]) -> None:
    """Emit blocks into a python-docx Document."""
    for block in blocks:
        paragraph = doc.add_paragraph(style=_DOCX_STYLES.get(block.kind))
        if block.kind == "number" and block.starts_list:
            _restart_numbering(doc, paragraph)
        for run in block.runs:
            r = paragraph.add_run(run.text)
            r.bold = run.bold or None
            r.italic = run.italic or None
            r.underline = run.underline or None


def docx_document(title: str, sections: list[tuple[str, list[Block]]]):
    """THE docx assembler — a title plus (heading, blocks) sections.

    Both the deterministic renderer and the web UI's export of hand-edited
    documents build through here, so the two can never format differently.
    """
    from docx import Document  # lazy: most commands never touch docx
    doc = Document()
    doc.add_heading(title, level=1)
    for heading, blocks in sections:
        doc.add_heading(heading, level=2)
        add_blocks(doc, blocks)
    return doc
